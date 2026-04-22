"""
Lesson 03 — File Translation Agent (Format Preserving)
========================================================
Claude orchestrates translation of an English .docx file
into Marathi and Hindi while preserving full formatting:
- Tables with all cells
- Paragraphs and headings
- Bold, italic, underline text styles
- Headers and footers

Output: Two properly formatted .docx files
        Token usage and cost report at the end
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog
from anthropic import Anthropic
from dotenv import load_dotenv
from deep_translator import GoogleTranslator
from docx import Document
from rich.console import Console
from rich.panel import Panel
from rich.table import Table

load_dotenv()

client = Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
console = Console()

# -----------------------------------------------------------------------
# Anthropic pricing for claude-sonnet-4-5 (per million tokens)
# -----------------------------------------------------------------------
PRICE_INPUT_PER_M  = 3.00   # $3.00 per million input tokens
PRICE_OUTPUT_PER_M = 15.00  # $15.00 per million output tokens

# -----------------------------------------------------------------------
# Token tracking
# -----------------------------------------------------------------------
total_input_tokens  = 0
total_output_tokens = 0


def update_tokens(input_tokens: int, output_tokens: int):
    global total_input_tokens, total_output_tokens
    total_input_tokens  += input_tokens
    total_output_tokens += output_tokens


def print_cost_report():
    """Print token usage and cost summary."""
    input_cost  = (total_input_tokens  / 1_000_000) * PRICE_INPUT_PER_M
    output_cost = (total_output_tokens / 1_000_000) * PRICE_OUTPUT_PER_M
    total_cost  = input_cost + output_cost

    table = Table(title="Token Usage & Cost Report", border_style="dim")
    table.add_column("Item",       style="dim", min_width=20)
    table.add_column("Tokens",     justify="right")
    table.add_column("Cost (USD)", justify="right")

    table.add_row(
        "Input tokens",
        f"{total_input_tokens:,}",
        f"${input_cost:.6f}"
    )
    table.add_row(
        "Output tokens",
        f"{total_output_tokens:,}",
        f"${output_cost:.6f}"
    )
    table.add_row(
        "[bold]Total[/bold]",
        f"[bold]{total_input_tokens + total_output_tokens:,}[/bold]",
        f"[bold]${total_cost:.6f}[/bold]"
    )

    console.print()
    console.print(table)
    console.print(
        f"[dim]Pricing: ${PRICE_INPUT_PER_M}/M input tokens, "
        f"${PRICE_OUTPUT_PER_M}/M output tokens (claude-sonnet-4-5)[/dim]\n"
    )


# -----------------------------------------------------------------------
# Translation helper
# -----------------------------------------------------------------------
def translate_string(text: str, target_language: str) -> str:
    """Translate a single string. Returns original if empty or error."""
    if not text or not text.strip():
        return text
    try:
        translator = GoogleTranslator(source="en", target=target_language)
        result = translator.translate(text)
        return result if result else text
    except Exception as e:
        console.print(f"  [red]Translation error: {e}[/red]")
        return text


# -----------------------------------------------------------------------
# Safe paragraph translator — handles None runs and paragraphs
# -----------------------------------------------------------------------
def translate_paragraphs(paragraphs, target_language: str) -> int:
    """Translate all runs in a list of paragraphs. Returns count of translated runs."""
    count = 0
    if not paragraphs:
        return count
    for para in paragraphs:
        try:
            if para is None:
                continue
            if not para.runs:
                continue
            for run in para.runs:
                if run is None:
                    continue
                if run.text and run.text.strip():
                    run.text = translate_string(run.text, target_language)
                    count += 1
        except Exception as e:
            console.print(f"  [dim red]Skipped paragraph: {e}[/dim red]")
            continue
    return count


# -----------------------------------------------------------------------
# Core: translate full .docx preserving all formatting
# -----------------------------------------------------------------------
def translate_docx(input_path: str, output_path: str, target_language: str, language_name: str) -> str:
    """
    Read input .docx, translate all text runs in-place, save output.
    Handles None paragraphs, runs, and complex document elements safely.
    """
    try:
        console.print(f"\n  [dim]Loading: {os.path.basename(input_path)}[/dim]")
        output_doc = Document(input_path)
        translated_count = 0

        # -- Paragraphs (outside tables) --
        translated_count += translate_paragraphs(output_doc.paragraphs, target_language)

        # -- Tables --
        for table in output_doc.tables:
            try:
                if table is None:
                    continue
                for row in table.rows:
                    if row is None:
                        continue
                    for cell in row.cells:
                        if cell is None:
                            continue
                        translated_count += translate_paragraphs(cell.paragraphs, target_language)
            except Exception as e:
                console.print(f"  [dim red]Skipped table: {e}[/dim red]")
                continue

        # -- Headers and footers --
        for section in output_doc.sections:
            try:
                if section is None:
                    continue
                if section.header:
                    translated_count += translate_paragraphs(section.header.paragraphs, target_language)
                if section.footer:
                    translated_count += translate_paragraphs(section.footer.paragraphs, target_language)
            except Exception as e:
                console.print(f"  [dim red]Skipped header/footer: {e}[/dim red]")
                continue

        output_doc.save(output_path)
        console.print(f"  [dim]Translated {translated_count} text runs to {language_name}[/dim]")
        console.print(f"  [green]Saved:[/green] {os.path.basename(output_path)}")
        return f"Successfully translated {translated_count} text runs and saved to {output_path}"

    except FileNotFoundError:
        return f"Error: File '{input_path}' not found"
    except Exception as e:
        return f"Error during translation: {str(e)}"


# -----------------------------------------------------------------------
# Tool definitions
# -----------------------------------------------------------------------
tools = [
    {
        "name": "translate_and_save",
        "description": (
            "Translates a .docx file from English to a target language "
            "preserving all formatting, tables, headers, footers and text styles. "
            "Saves the result as a new .docx file."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "input_path": {
                    "type": "string",
                    "description": "Full path of the English .docx file to translate"
                },
                "output_path": {
                    "type": "string",
                    "description": "Full path where the translated .docx file will be saved"
                },
                "target_language": {
                    "type": "string",
                    "description": "Language code: 'mr' for Marathi, 'hi' for Hindi"
                },
                "language_name": {
                    "type": "string",
                    "description": "Human readable language name e.g. Marathi, Hindi"
                }
            },
            "required": ["input_path", "output_path", "target_language", "language_name"]
        }
    }
]


# -----------------------------------------------------------------------
# Tool dispatcher
# -----------------------------------------------------------------------
def call_tool(tool_name: str, tool_input: dict) -> str:
    if tool_name == "translate_and_save":
        return translate_docx(
            input_path      = tool_input["input_path"],
            output_path     = tool_input["output_path"],
            target_language = tool_input["target_language"],
            language_name   = tool_input["language_name"]
        )
    return f"Unknown tool: {tool_name}"


# -----------------------------------------------------------------------
# Agent loop
# -----------------------------------------------------------------------
def run_agent(user_message: str):
    console.print(Panel(
        f"[bold]Task:[/bold] {user_message}",
        border_style="dim",
        expand=False
    ))

    messages = [{"role": "user", "content": user_message}]
    step = 0

    while True:
        step += 1
        console.print(f"\n[dim]--- Step {step} ---[/dim]")

        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=4096,
            tools=tools,
            messages=messages
        )

        update_tokens(response.usage.input_tokens, response.usage.output_tokens)
        console.print(
            f"[dim]Stop reason: {response.stop_reason} | "
            f"Tokens — in: {response.usage.input_tokens}, "
            f"out: {response.usage.output_tokens}[/dim]"
        )

        # Claude is done
        if response.stop_reason == "end_turn":
            for block in response.content:
                if hasattr(block, "text"):
                    console.print(f"\n[bold green]Claude:[/bold green] {block.text}")
            break

        # Claude wants to use tools
        if response.stop_reason == "tool_use":
            messages.append({
                "role": "assistant",
                "content": response.content
            })

            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    console.print(f"\n[cyan]Tool:[/cyan]  {block.name}")
                    console.print(f"[cyan]Input:[/cyan] language={block.input.get('language_name','?')}, "
                                  f"target={block.input.get('target_language','?')}")

                    result = call_tool(block.name, block.input)

                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": result
                    })

            messages.append({
                "role": "user",
                "content": tool_results
            })


# -----------------------------------------------------------------------
# Native Mac GUI file picker
# -----------------------------------------------------------------------
def browse_file() -> str:
    console.print(Panel(
        "[bold]File Translation Agent[/bold]\n"
        "[dim]Translates English .docx to Marathi and Hindi — formatting preserved[/dim]",
        border_style="dim",
        expand=False
    ))

    console.print("\n[dim]Opening file picker... (press Escape to cancel)[/dim]")

    root = tk.Tk()
    root.withdraw()
    root.call('wm', 'attributes', '.', '-topmost', True)

    file_path = filedialog.askopenfilename(
        title="Select English .docx file to translate",
        filetypes=[
            ("Word documents", "*.docx"),
            ("All files", "*.*")
        ]
    )

    root.destroy()

    if not file_path:
        console.print("[red]No file selected. Exiting.[/red]")
        sys.exit(0)

    console.print(f"[green]Selected:[/green] {os.path.basename(file_path)}\n")
    return file_path


# -----------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------
if __name__ == "__main__":
    if not os.environ.get("ANTHROPIC_API_KEY"):
        console.print("[red]ANTHROPIC_API_KEY not set in .env[/red]")
        sys.exit(1)

    # Step 1: pick file
    file_path = browse_file()

    # Step 2: build output paths
    output_dir = os.path.dirname(os.path.abspath(file_path))
    base_name  = os.path.splitext(os.path.basename(file_path))[0]

    marathi_output = os.path.join(output_dir, f"{base_name}_marathi.docx")
    hindi_output   = os.path.join(output_dir, f"{base_name}_hindi.docx")

    console.print(f"[dim]Output files:[/dim]")
    console.print(f"[dim]  Marathi → {base_name}_marathi.docx[/dim]")
    console.print(f"[dim]  Hindi   → {base_name}_hindi.docx[/dim]\n")

    # Step 3: run the agent
    run_agent(
        f"Please translate the file '{file_path}' to both Marathi and Hindi. "
        f"Save the Marathi version to '{marathi_output}' with language_name 'Marathi' "
        f"and language code 'mr'. "
        f"Save the Hindi version to '{hindi_output}' with language_name 'Hindi' "
        f"and language code 'hi'. "
        f"Preserve all formatting, tables, and document structure."
    )

    # Step 4: print token and cost report
    print_cost_report()